# -*- coding: utf-8 -*-
"""Insert functional module diagram into the opening report"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document(r'd:/ABS/司明-开题报告-基于SpringBoot的智能学习预警服务系统.docx')

# Find the paragraph containing "图1 系统功能结构图"
found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '图1' in text and '功能结构图' in text:
        found = True
        # Insert image BEFORE this caption paragraph
        # Add a new paragraph with the image just before this one
        img_para = para.insert_paragraph_before(para)
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = img_para.add_run()
        # Get image path - user provided as inline image, need to find it
        img_path = r'd:/ABS/image.fa3a2ce6a0.png'
        try:
            run.add_picture(img_path, width=Cm(16))
            print(f'Image inserted before: {text}')
        except Exception as e:
            print(f'Image error: {e}')
        
        # Also add spacing after image
        img_para.paragraph_format.space_after = Pt(6)
        break

if not found:
    print('Caption not found, trying alternative search...')
    for i, para in enumerate(doc.paragraphs):
        if '功能结构图' in para.text or '图1' in para.text:
            print(f'Found near line {i}: {para.text[:50]}')

# Save
output = r'd:/ABS/司明-开题报告-基于SpringBoot的智能学习预警服务系统.docx'
doc.save(output)
print(f'Saved to {output}')
